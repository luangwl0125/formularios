# cbcl_pais.py
import datetime
from pathlib import Path
import streamlit as st
from docx import Document
import smtplib
from email.message import EmailMessage
from io import BytesIO

st.set_page_config(page_title="Formulário CBCL", layout="centered")

# Página inicial
st.title("Formulário CBCL (Child Behavior Checklist) – Perguntas para o(a) Genitor(a)")
st.write("""
Este formulário é um instrumento padronizado de avaliação comportamental e tem como objetivo avaliar de forma sistemática e padronizada os comportamentos, habilidades sociais e sintomas emocionais da criança ou adolescente a partir da perspectiva dos cuidadores (geralmente pais ou responsáveis).
""")
st.write("""
**Atenção:** Os dados serão mantidos sob sigilo e utilizados apenas para fins clínicos, conforme o Código de Ética Profissional do Psicólogo (CFP) e a Resolução CFP nº 11/2018.  
**Psicólogo/Neuropsicólogo Responsável:** Luan Gama Wanderley Leite (CRP-15/3328)
""")

agree = st.checkbox("Concordo com o uso das informações para fins de avaliação psicológica.")
if agree:
    with st.form("cbcl_form"):
        # 1. Dados Pessoais
        st.markdown("### 1. Dados Pessoais")
        nome_responsavel     = st.text_input("Nome Completo")                  
        data_nascimento_str  = st.text_input("Data de Nascimento (DD/MM/AAAA)", placeholder="DD/MM/AAAA")
        nome_crianca         = st.text_input("Nome Completo da Criança ou Adolescente")
        parentesco           = st.text_input("Parentesco")

        # 2. Perguntas CBCL
        st.markdown("### 2. CBCL – Perguntas (Genitor(a))")
        perguntas = {
            1:  "Chora sem motivo aparente.",
            2:  "Demonstra tristeza prolongada.",
            3:  "Apresenta preocupação excessiva com coisas pequenas.",
            4:  "Queixa-se de medo sem razão clara.",
            5:  "Critica-se frequentemente.",
            6:  "Demonstra pensamento negativo sobre si mesmo.",
            7:  "Evita atividades de que costumava gostar.",
            8:  "Mostra inquietação interna (parece ‘tenso’).",
            9:  "Expressa sentimento de solidão.",
            10: "Tem dificuldade para dormir por preocupações.",
            11: "Evita contato social por insegurança.",
            12: "Queixa-se de falta de energia.",
            13: "Demonstra irritabilidade sem motivo forte.",
            14: "Apresenta baixa autoestima.",
            15: "Queixa-se de dores de cabeça frequentes.",
            16: "Reporta dores abdominais sem causa médica.",
            17: "Refere cansaço sem explicação.",
            18: "Tem náuseas ou enjoos frequentes.",
            19: "Queixa-se de dor no peito ou palpitações.",
            20: "Recurre a médicos sem necessidade real.",
            21: "Reclama de dores musculares.",
            22: "Sente tontura sem motivo orgânico.",
            23: "Apresenta alterações de apetite (muito ou pouco).",
            24: "Tem dificuldade de respiração sem causa.",
            25: "Reporta coceiras ou prurido sem motivo.",
            26: "Fala coisas que parecem sem sentido.",
            27: "Demonstra ideias fixas sem realidade.",
            28: "Tem dificuldade de manter linha de raciocínio.",
            29: "Fica repetindo palavras ou frases.",
            30: "Demonstra pensamentos confusos.",
            31: "Expressa crenças estranhas (ex.: alguém vai prejudicá-lo).",
            32: "Demonstra preocupação com ‘ser vigiado’.",
            33: "Diz que ouve sons ou vozes sem que existam.",
            34: "Tem dificuldade de distinguir fantasia e realidade.",
            35: "Demonstra ideias de grandeza sem base.",
            36: "Tem dificuldade em fazer amigos.",
            37: "Prefere ficar sozinho a brincar com colegas.",
            38: "Não compartilha brinquedos ou objetos.",
            39: "Tem dificuldade de cooperar em grupo.",
            40: "Demonstra insensibilidade ao que os outros sentem.",
            41: "É evitado por outras crianças.",
            42: "Ignora convites para atividades sociais.",
            43: "Não responde quando chamado pelo nome.",
            44: "Mostra falta de interesse em conversas.",
            45: "Não segue regras de jogos em grupo.",
            46: "Demonstra comportamento rude com amigos.",
            47: "Tem dificuldade em manter conversa.",
            48: "Reage mal a críticas de colegas.",
            49: "Desobedece a regras em casa.",
            50: "Ignora orientações dos professores.",
            51: "Rói unhas ou objetos de forma compulsiva.",
            52: "Torna-se teimoso em tarefas domésticas.",
            53: "Recusa-se a cumprir horários.",
            54: "Quebra acordos simples sem remorso.",
            55: "Rejeita limites de autoridade.",
            56: "Interrompe outras pessoas constantemente.",
            57: "Faz bagunça intencional em sala.",
            58: "Causa perturbação em locais públicos.",
            59: "Não aceita consequências de seus atos.",
            60: "Resiste a cooperar em tarefas em grupo.",
            61: "Fura fila propositalmente.",
            62: "Diminui a verdade para evitar problemas.",
            63: "Destrói objetos dos outros.",
            64: "Fala palavrões com frequência.",
            65: "Rouba objetos pequenos.",
            66: "Pratica vandalismo em casa ou escola.",
            67: "Engana adultos para benefício próprio.",
            68: "Agride fisicamente colegas.",
            69: "Faz bullying verbal.",
            70: "Participa de brigas sem provocação.",
            71: "Gasta dinheiro sem permissão.",
            72: "Usa mentiras para escapar de responsabilidades.",
            73: "Não consegue ficar parado.",
            74: "Fala demais.",
            75: "Tem dificuldade em esperar sua vez.",
            76: "Agita mãos ou pés constantemente.",
            77: "Muda de atividade sem concluir.",
            78: "Parece ‘impulsivo’ nas ações.",
            79: "Corre ou sobe em móveis sem razão.",
            80: "Tem dificuldade em controlar movimentos.",
            81: "Age sem pensar nas consequências.",
            82: "Interfere em atividades dos outros.",
            83: "Demonstra agressividade verbal.",
            84: "Tem explosões de raiva.",
            85: "Mostra ciúmes excessivo de irmãos.",
            86: "Tem medo de ir à escola.",
            87: "Não aceita mudanças de rotina.",
            88: "Tem pesadelos frequentes.",
            89: "Mostra comportamento desafiador com pais.",
            90: "Demonstra apego exagerado a objetos.",
            91: "Apresenta comportamentos ritualísticos.",
            92: "Tem dificuldades em expressar afeto.",
            93: "Apresenta bom desempenho escolar.",
            94: "Entrega tarefas no prazo.",
            95: "Participa de atividades extracurriculares.",
            96: "Ajuda em atividades domésticas.",
            97: "Interage bem com familiares.",
            98: "Mostra interesse em hobbies.",
            99: "Trabalha bem em grupo.",
            100:"Mostra iniciativa em tarefas.",
            101:"Respeita regras de convivência.",
            102:"Demonstra responsabilidade com pertences.",
            103:"Participa de esportes com prazer.",
            104:"Demonstra concentração em atividades.",
            105:"Lê ou estuda por conta própria.",
            106:"Mostra criatividade em brincadeiras.",
            107:"Ajuda colegas com dificuldades.",
            108:"Expressa opiniões de forma coerente.",
            109:"Desenvolve novas amizades com facilidade.",
            110:"Mantém rotina de sono adequada.",
            111:"Alimenta-se de forma saudável.",
            112:"Demonstra autocontrole em situações de estresse.",
            113:"Cumpre horários de estudo e lazer."
        }
        respostas = {n: st.radio(f"{n}. {t}", ["Sim", "Não"]) for n, t in perguntas.items()}
        observacoes = st.text_area("Você gostaria de acrescentar alguma observação?")

        submitted = st.form_submit_button("Enviar")
        if submitted:
            # validação da data
            try:
                data_nascimento = datetime.datetime.strptime(data_nascimento_str, "%d/%m/%Y").date()
            except ValueError:
                st.error("Data de Nascimento inválida. Use o formato DD/MM/AAAA.")
            else:
                # feedback de progresso
                prog = st.progress(0)
                st.info("Gerando documento…")
                # Geração de Word (.docx) – formato mais estável
                doc = Document()
                doc.add_heading("Formulário CBCL - Respostas", level=1)
                doc.add_paragraph(f"Responsável: {nome_responsavel}")
                doc.add_paragraph(f"Data de Nascimento: {data_nascimento_str}")
                doc.add_paragraph(f"Criança/Adolescente: {nome_crianca}")
                doc.add_paragraph(f"Parentesco: {parentesco}")
                doc.add_paragraph("")
                prog.progress(30)
                for n, txt in perguntas.items():
                    doc.add_paragraph(f"{n}. {txt} Resposta: {respostas[n]}")
                prog.progress(60)
                if observacoes.strip():
                    doc.add_paragraph("")
                    doc.add_heading("Observações", level=2)
                    doc.add_paragraph(observacoes)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                prog.progress(80)

                # envio por e-mail
                msg = EmailMessage()
                msg["Subject"] = "CBCL - Respostas do Formulário"
                msg["From"]    = st.secrets["email"]["sender"]
                msg["To"]      = st.secrets["email"]["recipient"]
                msg.set_content("Segue em anexo o arquivo .docx com as respostas do formulário CBCL.")
                msg.add_attachment(
                    buffer.read(),
                    maintype="application",
                    subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
                    filename="cbcl_respostas.docx"
                )
                with smtplib.SMTP(st.secrets["smtp"]["server"], st.secrets["smtp"]["port"]) as server:
                    server.starttls()
                    server.login(st.secrets["email"]["sender"], st.secrets["email"]["password"])
                    server.send_message(msg)
                prog.progress(100)
                st.success("Arquivo enviado com sucesso!")
